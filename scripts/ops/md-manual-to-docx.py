# md-manual-to-docx.py
# Merge the HT160S manual Markdown files (cover + chapters + appendices) into one
# Traditional-Chinese .docx with a TOC field, heading hierarchy, embedded
# screenshots and tables. Uses python-docx only (no pandoc required).
#
# Usage: python scripts/ops/md-manual-to-docx.py <manual_dir> <out.docx>
import sys, os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MANUAL = sys.argv[1] if len(sys.argv) > 1 else "D:/HT160S_BCB/docs/manual"
OUT    = sys.argv[2] if len(sys.argv) > 2 else os.path.join(MANUAL, "HT160S_操作手冊.docx")
CJK    = "Microsoft JhengHei"
CONTENT_WIDTH_IN = 6.3  # inches available between 1" margins on US Letter

FILES = ["00-cover.md"] + ["%02d-%s.md" % (i, n) for i, n in [
    (1,"safety"),(2,"overview"),(3,"panel-startup"),(4,"main-screen"),(5,"maintenance"),
    (6,"config"),(7,"teach"),(8,"offset"),(9,"speed"),(10,"io"),(11,"motor-test"),
    (12,"secs-amr"),(13,"alarms"),(14,"module-flows"),(15,"lotbin-mode"),(16,"faq")]] + \
    ["A1-field-checklist.md","B1-io-table.md","C1-motor-table.md","D1-alarm-list.md"]

def set_cjk(style):
    rpr = style.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), CJK); rf.set(qn('w:ascii'), CJK); rf.set(qn('w:hAnsi'), CJK)

doc = Document()
# default + heading fonts -> CJK
for nm in ["Normal","Title","Heading 1","Heading 2","Heading 3","Heading 4","List Bullet","List Number"]:
    try: set_cjk(doc.styles[nm])
    except KeyError: pass
doc.styles["Normal"].font.size = Pt(10.5)

def add_toc(d):
    p = d.add_paragraph()
    run = p.add_run()
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u')
    t = OxmlElement('w:t'); t.text = "（開啟後請按 Ctrl+A 再按 F9 更新目錄）"
    r = OxmlElement('w:r'); r.append(t); fld.append(r); p._p.append(fld)

def inline_runs(p, text):
    # handle **bold**; strip backticks and link syntax
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # [t](u)->t
    text = text.replace('`', '')
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for seg in parts:
        if not seg: continue
        if seg.startswith('**') and seg.endswith('**'):
            r = p.add_run(seg[2:-2]); r.bold = True
        else:
            p.add_run(seg)

def emit_table(rows):
    # rows: list of cell-lists; first row header, second was separator (already removed)
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    for ri, cells in enumerate(rows):
        cdx = t.add_row().cells
        for ci in range(ncol):
            val = cells[ci] if ci < len(cells) else ""
            cell = cdx[ci]; cell.text = ""
            para = cell.paragraphs[0]
            inline_runs(para, val.strip())
            for run in para.runs:
                run.font.size = Pt(8 if ncol >= 8 else 9)
                if ri == 0: run.bold = True
    doc.add_paragraph()

img_re = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
def process_file(path, first):
    if not first:
        doc.add_page_break()
    lines = open(path, encoding='utf-8').read().split('\n')
    i = 0; in_code = False; code_buf = []
    while i < len(lines):
        ln = lines[i]
        s = ln.rstrip()
        if s.strip().startswith('```'):
            if in_code:
                cp = doc.add_paragraph(); cp.paragraph_format.left_indent = Inches(0.2)
                r = cp.add_run('\n'.join(code_buf)); r.font.name = 'Consolas'; r.font.size = Pt(9)
                code_buf = []; in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(ln); i += 1; continue
        # table block
        if s.startswith('|') and i+1 < len(lines) and re.match(r'^\s*\|[\s:\-|]+\|\s*$', lines[i+1]):
            block = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                block.append(lines[i]); i += 1
            rows = []
            for bi, brow in enumerate(block):
                if bi == 1:  # separator
                    continue
                cells = [c for c in brow.strip().strip('|').split('|')]
                rows.append(cells)
            emit_table(rows)
            continue
        # image
        m = img_re.search(s)
        if m:
            alt, src = m.group(1), m.group(2)
            ip = src if os.path.isabs(src) else os.path.join(MANUAL, src)
            if os.path.exists(ip):
                try:
                    doc.add_picture(ip, width=Inches(CONTENT_WIDTH_IN))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph("[image: %s (%s)]" % (src, e))
            else:
                doc.add_paragraph("[缺圖版位: %s]" % src)
            i += 1; continue
        # headings
        hm = re.match(r'^(#{1,4})\s+(.*)$', s)
        if hm:
            lvl = len(hm.group(1)); doc.add_heading(hm.group(2).strip(), level=min(lvl,4)); i += 1; continue
        # hr
        if re.match(r'^\s*---+\s*$', s):
            i += 1; continue
        # blockquote
        if s.lstrip().startswith('>'):
            qp = doc.add_paragraph(); qp.paragraph_format.left_indent = Inches(0.25)
            txt = s.lstrip()[1:].strip()
            inline_runs(qp, txt)
            for r in qp.runs: r.italic = True; r.font.color.rgb = RGBColor(0x55,0x55,0x66)
            i += 1; continue
        # lists
        lm = re.match(r'^(\s*)(\d+)\.\s+(.*)$', ln)
        if lm:
            p = doc.add_paragraph(style='List Number'); inline_runs(p, lm.group(3)); i += 1; continue
        lm = re.match(r'^(\s*)[-*]\s+(.*)$', ln)
        if lm:
            p = doc.add_paragraph(style='List Bullet'); inline_runs(p, lm.group(2)); i += 1; continue
        # blank
        if not s.strip():
            i += 1; continue
        # plain paragraph
        p = doc.add_paragraph(); inline_runs(p, s); i += 1

# Title page + TOC
doc.add_heading("HT160S 操作手冊（工程版）", level=0)
doc.add_paragraph("由 .claude/workflows/ht160s-operation-manual.js 自原始碼產生；草稿，待現場校對。")
doc.add_heading("目錄", level=1)
add_toc(doc)

for idx, fn in enumerate(FILES):
    p = os.path.join(MANUAL, fn)
    if os.path.exists(p):
        process_file(p, first=False)
    else:
        print("MISSING:", fn)

doc.save(OUT)
print("saved", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
